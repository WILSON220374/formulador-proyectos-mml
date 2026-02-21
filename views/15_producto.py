import streamlit as st
import os
import io
from datetime import datetime
from session_state import inicializar_session

# --- IMPORTACIÓN DE LIBRERÍAS (WORD Y PDF) ---
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("⚠️ Falta la librería para Word. Agrega 'python-docx' a tu requirements.txt")
    st.stop()

try:
    from fpdf import FPDF
except ImportError:
    st.error("⚠️ Falta la librería para PDF. Agrega 'fpdf2' a tu requirements.txt")
    st.stop()

# 1. Asegurar persistencia 
inicializar_session()

# --- DISEÑO PROFESIONAL (CSS) ---
st.markdown("""
    <style>
    .block-container { padding-bottom: 12rem !important; }
    .titulo-seccion { font-size: 30px !important; font-weight: 800 !important; color: #1E3A8A; margin-bottom: 5px; }
    .subtitulo-gris { font-size: 16px !important; color: #666; margin-bottom: 15px; }
    .header-tabla { font-weight: 800; color: #1E3A8A; margin-bottom: 10px; font-size: 1.1rem; text-transform: uppercase; border-bottom: 2px solid #1E3A8A; padding-bottom: 5px;}
    .readonly-box { border: 1px solid #d1d5db; border-radius: 8px; padding: 12px; background-color: #f3f4f6; color: #1E3A8A; font-weight: 800; text-align: center; font-size: 1.2rem;}
    </style>
""", unsafe_allow_html=True)

# --- ENCABEZADO ---
col_t, col_img = st.columns([4, 1], vertical_alignment="center")
with col_t:
    st.markdown('<div class="titulo-seccion">📄 16. Generador de Reportes</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitulo-gris">Configuración de portada y exportación de documentos.</div>', unsafe_allow_html=True)
with col_img:
    if os.path.exists("unnamed.jpg"):
        st.image("unnamed.jpg", use_container_width=True)

st.divider()

# ==========================================
# 📘 1. CONFIGURACIÓN DE LA PORTADA
# ==========================================
st.markdown('<div class="header-tabla">📘 1. Configuración de la Portada</div>', unsafe_allow_html=True)

# 1. Traer el Nombre del Proyecto de la Hoja 15
nombre_proyecto = st.session_state.get('nombre_proyecto_libre', 'AÚN NO SE HA DEFINIDO EL NOMBRE DEL PROYECTO (Vaya a la Hoja 15)')

st.write("**Nombre del Proyecto:**")
st.markdown(f'<div class="readonly-box">{nombre_proyecto.upper()}</div><br>', unsafe_allow_html=True)

# 2. Carga de Imágenes (CON VISTA PREVIA)
col_img1, col_img2 = st.columns(2)
with col_img1:
    st.info("🖼️ **Logo de la Entidad** (Irá en la esquina superior derecha)")
    logo_entidad = st.file_uploader("Sube el logo", type=["png", "jpg", "jpeg"], key="logo_portada")
    # Mostrar la imagen si el usuario subió algo
    if logo_entidad is not None:
        st.image(logo_entidad, width=150, caption="Vista previa del Logo")

with col_img2:
    st.info("📸 **Imagen Central** (Irá en el centro de la portada)")
    img_portada = st.file_uploader("Sube la imagen central", type=["png", "jpg", "jpeg"], key="img_portada")
    # Mostrar la imagen si el usuario subió algo
    if img_portada is not None:
        st.image(img_portada, width=250, caption="Vista previa de la Imagen Central")

st.write("") # Espacio

# 3. Datos a digitar
col_d1, col_d2, col_d3 = st.columns([2, 2, 1])
with col_d1:
    entidad_formulo = st.text_input("Entidad que formula el proyecto", placeholder="Ej: Alcaldía de Tunja")
with col_d2:
    lugar_presentacion = st.text_input("Lugar de presentación", placeholder="Ej: Tunja, Boyacá")
with col_d3:
    anio_presentacion = st.text_input("Año", value="2026")

st.divider()

# ==========================================
# 📑 2. MENÚ DE SELECCIÓN DE CONTENIDO
# ==========================================
st.markdown('<div class="header-tabla">📑 2. Selección de Contenido</div>', unsafe_allow_html=True)

with st.container(border=True):
    st.markdown("**Hoja: Diagnóstico (Árbol de Problemas)**")
    chk_problema = st.checkbox("El Problema Central", value=True)
    chk_sintomas = st.checkbox("Síntomas (Efectos)", value=True)
    chk_causas = st.checkbox("Causas Inmediatas", value=True)

st.divider()

# ==========================================
# 🛑 TEXTOS DE PRUEBA INTERNOS
# ==========================================
texto_prob_prueba = "Texto de prueba: La alta tasa de accidentalidad en la vía principal debido a la falta de mantenimiento."
texto_sintomas_prueba = "Texto de prueba: 1. Incremento en tiempos de traslado. 2. Daños constantes a los vehículos."
texto_causas_prueba = "Texto de prueba: 1. Deterioro de la capa asfáltica. 2. Ausencia de mantenimiento."

# ==========================================
# ⚙️ MOTOR DE GENERACIÓN WORD (Temporal sin imágenes en el papel)
# ==========================================
def generar_word():
    doc = Document()
    titulo = doc.add_heading("Reporte de Formulación de Proyecto", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    
    doc.add_page_break()
    doc.add_heading("1. Diagnóstico y Problema", level=1)
    
    if chk_problema:
        doc.add_heading("1.1 El Problema Central", level=2)
        doc.add_paragraph(texto_prob_prueba)
    if chk_sintomas:
        doc.add_heading("1.2 Síntomas (Efectos)", level=2)
        doc.add_paragraph(texto_sintomas_prueba)
    if chk_causas:
        doc.add_heading("1.3 Causas Inmediatas", level=2)
        doc.add_paragraph(texto_causas_prueba)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# ⚙️ MOTOR DE GENERACIÓN PDF (Temporal sin imágenes en el papel)
# ==========================================
def generar_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "Reporte de Formulacion de Proyecto", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.add_page()
    pdf.set_font("helvetica", "B", 14)
    pdf.cell(0, 10, "1. Diagnostico y Problema", new_x="LMARGIN", new_y="NEXT")
    return pdf.output()

# --- 3. BOTONES DE DESCARGA ---
st.markdown('<div class="header-tabla">📥 3.
