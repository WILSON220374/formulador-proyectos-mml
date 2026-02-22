import streamlit as st
import os
import io
import pandas as pd
import requests
from session_state import inicializar_session

# --- IMPORTACIÓN DE LIBRERÍAS (WORD, PDF Y GRAPHVIZ) ---
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    st.error("⚠️ Falta la librería para Word. Agrega 'python-docx' a tu requirements.txt")
    st.stop()

try:
    from fpdf import FPDF
except ImportError:
    st.error("⚠️ Falta la librería para PDF. Agrega 'fpdf2' a tu requirements.txt")
    st.stop()

try:
    import graphviz
except ImportError:
    st.error("⚠️ Falta la librería Graphviz. Agrega 'graphviz' a tu requirements.txt")
    st.stop()

# 1. Asegurar persistencia 
inicializar_session()

# --- DISEÑO PROFESIONAL (CSS) ---
st.markdown("""
    <style>
    .block-container { padding-bottom: 12rem !important; }
    .titulo-seccion { font-size: 30px !important; font-weight: 800 !important; color: #1E3A8A; margin-bottom: 5px; }
    .subtitulo-gris { font-size: 16px !important; color: #666; margin-bottom: 15px; }
    .header-tabla { font-weight: 800; color: #1E3A8A; margin-bottom: 10px; font-size: 1.1rem; text-transform: uppercase; border-bottom: 2px solid #1E3A8A; padding-bottom: 5px;}
    .readonly-box { border: 1px solid #d1d5db; border-radius: 8px; padding: 12px; background-color: #f3f4f6; color: #1E3A8A; font-weight: 800; text-align: center; font-size: 1.2rem;}
    .readonly-autores { border: 1px solid #d1d5db; border-radius: 8px; padding: 10px; background-color: #f3f4f6; color: #374151; font-weight: 600; text-align: center; font-size: 1rem;}
    </style>
""", unsafe_allow_html=True)

# --- ENCABEZADO ---
col_t, col_img = st.columns([4, 1], vertical_alignment="center")
with col_t:
    st.markdown('<div class="titulo-seccion">📄 16. Generador de Reportes</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitulo-gris">Configuración de portada y exportación del documento final.</div>', unsafe_allow_html=True)
with col_img:
    if os.path.exists("unnamed.jpg"):
        st.image("unnamed.jpg", use_container_width=True)

st.divider()

# ==========================================
# 📘 1. CONFIGURACIÓN DE LA PORTADA
# ==========================================
st.markdown('<div class="header-tabla">📘 1. Configuración de la Portada</div>', unsafe_allow_html=True)

nombre_proyecto = st.session_state.get('nombre_proyecto_libre', 'NOMBRE DEL PROYECTO NO DEFINIDO') 

st.write("**Nombre del Proyecto:**")
st.markdown(f'<div class="readonly-box">{nombre_proyecto.upper()}</div><br>', unsafe_allow_html=True)

col_up1, col_up2 = st.columns(2)
with col_up1:
    st.info("🖼️ **Logo de la Entidad** (Irá en el encabezado de todas las páginas)")
    logo_entidad = st.file_uploader("Sube el logo", type=["png", "jpg", "jpeg"], key="logo_portada")

with col_up2:
    st.info("📸 **Imagen Central** (Irá en el centro de la portada)")
    img_portada = st.file_uploader("Sube la imagen central", type=["png", "jpg", "jpeg"], key="img_portada")

if logo_entidad is not None or img_portada is not None:
    col_prev1, col_prev2 = st.columns(2)
    with col_prev1:
        if logo_entidad is not None:
            st.image(logo_entidad, width=150, caption="Logo listo para el encabezado")
    with col_prev2:
        if img_portada is not None:
            st.image(img_portada, width=300, caption="Imagen Central lista")

st.write("") 

nombres_formuladores = "No se encontraron formuladores registrados en la Hoja 1"
nombres_display = nombres_formuladores

if "df_equipo" in st.session_state and isinstance(st.session_state["df_equipo"], pd.DataFrame):
    df = st.session_state["df_equipo"]
    if "Nombre" in df.columns:
        nombres_lista = df["Nombre"].dropna().astype(str).tolist()
        nombres_validos = [n for n in nombres_lista if n.strip() != ""]
        if nombres_validos:
            nombres_formuladores = "\n".join(nombres_validos) 
            nombres_display = ", ".join(nombres_validos) 

if nombres_formuladores.startswith("No se"):
    integrantes = st.session_state.get("integrantes", [])
    if isinstance(integrantes, list):
        nombres_validos = []
        for p in integrantes:
            if isinstance(p, dict):
                nombre = str(p.get("Nombre Completo", p.get("Nombre", p.get("nombre", "")))).strip()
                if nombre:
                    nombres_validos.append(nombre)
        if nombres_validos:
            nombres_formuladores = "\n".join(nombres_validos)
            nombres_display = ", ".join(nombres_validos)

st.write("**Presentado por (Equipo Formulador):**")
st.markdown(f'<div class="readonly-autores">{nombres_display}</div><br>', unsafe_allow_html=True)

col_d1, col_d2 = st.columns(2)
with col_d1:
    entidad_formulo = st.text_input("Entidad que formula el proyecto", placeholder="Ej: Alcaldía de Tunja")
with col_d2:
    division = st.text_input("División / Dependencia", placeholder="Ej: Secretaría de Infraestructura")

col_d3, col_d4 = st.columns(2)
with col_d3:
    lugar_presentacion = st.text_input("Lugar de presentación", value="Tunja, Boyacá")
with col_d4:
    anio_presentacion = st.text_input("Año", value="2026")

st.divider()

# ==========================================
# 📑 2. SELECCIÓN Y DILIGENCIAMIENTO DE CONTENIDO
# ==========================================
st.markdown('<div class="header-tabla">📑 2. Configuración del Contenido</div>', unsafe_allow_html=True)
st.write("Diligencia las siguientes secciones que se incluirán al inicio de tu documento:")

instrucciones_resumen = "El resumen es el elemento fundamental para dar contexto sobre el proyecto, en este sentido escriba en máximo 15 líneas de manera clara, sencilla, directa y concisa el resumen del contenido del proyecto, que permitan dar una idea sobre el alcance, componentes y productos esperados."
texto_resumen = st.text_area("2. RESUMEN DEL PROYECTO", placeholder=instrucciones_resumen, height=150)

texto_normativo = st.text_area("3. MARCO NORMATIVO", placeholder="Escriba aquí el marco normativo del proyecto...", height=150)

st.divider()

# ==========================================
# 📥 EXTRACCIÓN DE DATOS DE LA MEMORIA
# ==========================================

# 1. Plan de Desarrollo
_plan_nombre = str(st.session_state.get('plan_nombre', '')).strip()
_plan_eje = str(st.session_state.get('plan_eje', '')).strip()
_plan_programa = str(st.session_state.get('plan_programa', '')).strip()

if any([_plan_nombre, _plan_eje, _plan_programa]):
    plan_desarrollo = (
        f"Nombre del Plan: {_plan_nombre if _plan_nombre else 'No definido'}\n"
        f"Eje: {_plan_eje if _plan_eje else 'No definido'}\n"
        f"Programa: {_plan_programa if _plan_programa else 'No definido'}"
    )
else:
    plan_desarrollo = st.session_state.get('plan_desarrollo', 'No se ha registrado información en la Hoja 15.')

# 2. Justificación
just_1 = str(st.session_state.get('temp_justificacion', '')).strip()
just_2 = str(st.session_state.get('justificacion_arbol_objetivos_final', '')).strip()
just_3 = str(st.session_state.get('arbol_objetivos_final', {}).get('referencia_manual', {}).get('justificacion', '')).strip()
just_4 = str(st.session_state.get('justificacion_proyecto', '')).strip()

justificacion = just_1 if just_1 else (just_2 if just_2 else (just_3 if just_3 else just_4))
if not justificacion:
    justificacion = 'No se ha registrado información en la Hoja 7.'

# 3. Datos de Localización (Hoja 9)
zona_data = st.session_state.get('descripcion_zona', {})
ruta_mapa = zona_data.get('ruta_mapa')
ruta_foto1 = zona_data.get('ruta_foto1')
ruta_foto2 = zona_data.get('ruta_foto2')


# --- 4. EXTRACCIÓN DE LA SECCIÓN 6: PROBLEMA (HOJA 8 Y 10) ---

# Funciones auxiliares para parsear diccionarios de la Hoja 10
def _a_texto_dict(item):
    if isinstance(item, dict): return item
    if item is None: return None
    return {"texto": str(item)}

def _a_lista_dicts(valor):
    if valor is None: return []
    if isinstance(valor, list):
        out = []
        for it in valor:
            d = _a_texto_dict(it)
            if d is not None: out.append(d)
        return out
    if isinstance(valor, dict): return [valor]
    return [{"texto": str(valor)}]

desc_prob_data = st.session_state.get('descripcion_problema', {})
narrativa_problema = desc_prob_data.get('redaccion_narrativa', 'No se ha registrado la descripción detallada.')
if not narrativa_problema.strip(): narrativa_problema = 'No se ha registrado la descripción detallada.'

antecedentes_problema = desc_prob_data.get('antecedentes', 'No se han registrado antecedentes.')
if not antecedentes_problema.strip(): antecedentes_problema = 'No se han registrado antecedentes.'

# Reconstruir Tabla de Magnitud
tabla_datos_prob = desc_prob_data.get('tabla_datos', {})
datos_h8 = st.session_state.get('arbol_problemas_final', {})
if not isinstance(datos_h8, dict): datos_h8 = {}

pp_list = _a_lista_dicts(datos_h8.get("Problema Principal"))
pc_txt = pp_list[0].get("texto", "") if pp_list and isinstance(pp_list[0], dict) else ""

lista_causas = [c.get("texto") for c in _a_lista_dicts(datos_h8.get("Causas Directas")) + _a_lista_dicts(datos_h8.get("Causas Indirectas")) if isinstance(c, dict) and c.get("texto")]
lista_efectos = [e.get("texto") for e in _a_lista_dicts(datos_h8.get("Efectos Directos")) + _a_lista_dicts(datos_h8.get("Efectos Indirectos")) if isinstance(e, dict) and e.get("texto")]

filas_magnitud = []
if pc_txt:
    filas_magnitud.append({
        "Categoría": "PROBLEMA CENTRAL",
        "Descripción": pc_txt,
        "Magnitud": tabla_datos_prob.get("m_pc", ""),
        "Unidad": tabla_datos_prob.get("u_pc", ""),
        "Cantidad": tabla_datos_prob.get("c_pc", "")
    })
for i, txt in enumerate(lista_causas):
    filas_magnitud.append({
        "Categoría": f"CAUSA {i+1}",
        "Descripción": txt,
        "Magnitud": tabla_datos_prob.get(f"m_causa_{i}", ""),
        "Unidad": tabla_datos_prob.get(f"u_causa_{i}", ""),
        "Cantidad": tabla_datos_prob.get(f"c_causa_{i}", "")
    })
for i, txt in enumerate(lista_efectos):
    filas_magnitud.append({
        "Categoría": f"EFECTO {i+1}",
        "Descripción": txt,
        "Magnitud": tabla_datos_prob.get(f"m_efecto_{i}", ""),
        "Unidad": tabla_datos_prob.get(f"u_efecto_{i}", ""),
        "Cantidad": tabla_datos_prob.get(f"c_efecto_{i}", "")
    })

df_magnitud_reconstruida = pd.DataFrame(filas_magnitud)

# 5. Población
df_poblacion_general = st.session_state.get('df_poblacion_general', None)
df_pob_sexo = st.session_state.get('df_pob_sexo', None)
df_pob_edad = st.session_state.get('df_pob_edad', None)
analisis_poblacion = st.session_state.get('analisis_poblacion', 'No se ha registrado análisis de población.')

# 6. Participantes
df_matriz_interesados = st.session_state.get('df_matriz_interesados', None)
df_mapa_influencia = st.session_state.get('df_mapa_influencia', None)
df_analisis_participantes = st.session_state.get('df_analisis_participantes', None)

# 7. Objetivos
arbol_objetivos_img = st.session_state.get('arbol_objetivos_img', None)
objetivo_general = st.session_state.get('objetivo_general', 'No se ha definido el objetivo general.')
objetivos_especificos = st.session_state.get('objetivos_especificos_lista', 'No se han definido objetivos específicos.')

# 8. Alternativas
alternativas_consolidadas = st.session_state.get('alternativas_consolidadas', 'No se han registrado alternativas consolidadas.')
df_evaluacion_alt = st.session_state.get('df_evaluacion_alt', None)
alternativa_seleccionada = st.session_state.get('alternativa_seleccionada', 'No se ha seleccionado ninguna alternativa.')


# ==========================================
# ⚙️ MOTORES AUXILIARES Y WORD
# ==========================================

def redibujar_arbol_problemas(arbol_data):
    """Motor de respaldo para dibujar el árbol de problemas usando Graphviz en caso de que la foto en memoria se borre"""
    try:
        dot = graphviz.Digraph(format='png')
        dot.attr(rankdir='BT')
        dot.attr('node', shape='box', style='filled', fontname='Helvetica', margin='0.2')

        pp_list = _a_lista_dicts(arbol_data.get("Problema Principal"))
        pc_txt = pp_list[0].get("texto", "Problema Central") if pp_list else "Problema Central"
        dot.node('PC', str(pc_txt), fillcolor='#FCA5A5')

        causas = _a_lista_dicts(arbol_data.get("Causas Directas")) + _a_lista_dicts(arbol_data.get("Causas Indirectas"))
        for i, ca in enumerate(causas):
            dot.node(f'C_{i}', ca.get('texto', ''), fillcolor='#FEF3C7')
            dot.edge(f'C_{i}', 'PC')

        efectos = _a_lista_dicts(arbol_data.get("Efectos Directos")) + _a_lista_dicts(arbol_data.get("Efectos Indirectos"))
        for i, ef in enumerate(efectos):
            dot.node(f'E_{i}', ef.get('texto', ''), fillcolor='#DBEAFE')
            dot.edge('PC', f'E_{i}')

        return io.BytesIO(dot.pipe())
    except Exception as e:
        return None

def agregar_tabla_word(doc, df):
    if isinstance(df, pd.DataFrame) and not df.empty:
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Table Grid'
        hdr_cells = t.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for index, row in df.iterrows():
            row_cells = t.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
    else:
        p = doc.add_paragraph()
        p.add_run("No se registraron datos en esta tabla.").italic = True

def descargar_y_pegar_imagen(doc, url, ancho):
    if url:
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(io.BytesIO(response.content), width=Inches(ancho))
                return True
        except Exception as e:
            pass
    return False

def generar_word():
    doc = Document()
    
    # --- CONFIGURACIÓN DE ENCABEZADO Y PIE DE PÁGINA ---
    section = doc.sections[0]
    section.different_first_page_header_footer = True 
    
    def crear_encabezado(hdr_obj):
        htable = hdr_obj.add_table(rows=1, cols=2, width=Inches(6.5))
        htable.autofit = False
        htable.columns[0].width = Inches(5.0) 
        htable.columns[1].width = Inches(1.5) 
        h_izq = htable.cell(0, 0).paragraphs[0]
        h_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_izq.add_run(nombre_proyecto.upper()).bold = True
        h_der = htable.cell(0, 1).paragraphs[0]
        h_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if logo_entidad is not None:
            logo_entidad.seek(0)
            h_der.add_run().add_picture(io.BytesIO(logo_entidad.getvalue()), width=Inches(0.6))
        p_line = hdr_obj.add_paragraph()
        pPr = p_line._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    crear_encabezado(section.first_page_header)
    crear_encabezado(section.header)
    
    footer = section.footer
    p_line_f = footer.paragraphs[0]
    pPr_f = p_line_f._p.get_or_add_pPr()
    pBdr_f = OxmlElement('w:pBdr')
    bottom_f = OxmlElement('w:bottom')
    bottom_f.set(qn('w:val'), 'single')
    bottom_f.set(qn('w:sz'), '6') 
    bottom_f.set(qn('w:space'), '1')
    bottom_f.set(qn('w:color'), 'auto')
    pBdr_f.append(bottom_f)
    pPr_f.append(pBdr_f)
    
    ftable = footer.add_table(rows=1, cols=3, width=Inches(7.0))
    ftable.autofit = False
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    ftable._tbl.tblPr.append(tblLayout)
    
    ancho_izq = Inches(0.5)
    ancho_cen = Inches(6.0) 
    ancho_der = Inches(0.5)
    
    ftable.columns[0].width = ancho_izq
    ftable.columns[1].width = ancho_cen
    ftable.columns[2].width = ancho_der
    for row in ftable.rows:
        row.cells[0].width = ancho_izq
        row.cells[1].width = ancho_cen
        row.cells[2].width = ancho_der
    
    f_centro = ftable.cell(0, 1).paragraphs[0]
    f_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_centro.paragraph_format.space_after = Pt(0) 
    f_centro.paragraph_format.space_before = Pt(0)
    
    texto_entidad = entidad_formulo if entidad_formulo else ""
    texto_div = division if division else ""
    
    if texto_entidad:
        r_centro1 = f_centro.add_run(texto_entidad.upper())
        r_centro1.font.size = Pt(9)
        r_centro1.italic = True
        r_centro1.font.color.rgb = RGBColor(128, 128, 128)
        
    if texto_div:
        if texto_entidad:
            f_centro2 = ftable.cell(0, 1).add_paragraph()
        else:
            f_centro2 = f_centro
            
        f_centro2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_centro2.paragraph_format.space_after = Pt(0)
        f_centro2.paragraph_format.space_before = Pt(0)
        
        r_centro2 = f_centro2.add_run(texto_div.upper())
        r_centro2.font.size = Pt(9)
        r_centro2.italic = True
        r_centro2.font.color.rgb = RGBColor(128, 128, 128)
    
    f_der = ftable.cell(0, 2).paragraphs[0]
    f_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_num = f_der.add_run()
    r_num.font.size = Pt(10)
    r_num.font.color.rgb = RGBColor(128, 128, 128)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_num._r.append(fldChar1)
    r_num._r.append(instrText)
    r_num._r.append(fldChar2)
    r_num._r.append(fldChar3)

    # --- 1. PORTADA ---
    doc.add_paragraph("\n\n") 
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titulo = p_titulo.add_run(nombre_proyecto.upper())
    r_titulo.bold = True
    r_titulo.font.size = Pt(20)
    
    doc.add_paragraph("\n")
    if img_portada is not None:
        img_portada.seek(0)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(img_portada.getvalue()), width=Inches(3.8))
        
    doc.add_paragraph("\n")
    if entidad_formulo:
        doc.add_paragraph(entidad_formulo.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if division:
        doc.add_paragraph(division.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph("\n")
    p_presentado = doc.add_paragraph()
    p_presentado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_presentado.add_run("Presentado por:\n").italic = True
    p_presentado.add_run(nombres_formuladores).bold = True
    
    doc.add_paragraph("\n")
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto_lugar = lugar_presentacion if lugar_presentacion else ""
    texto_anio = anio_presentacion if anio_presentacion else ""
    p_pie.add_run(f"{texto_lugar}\n{texto_anio}".strip()).bold = True
    
    doc.add_page_break()
    
    # --- 2. INICIO DEL CONTENIDO ---
    p_tit_cont = doc.add_paragraph()
    p_tit_cont.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t_cont = p_tit_cont.add_run(nombre_proyecto.upper())
    r_t_cont.bold = True
    r_t_cont.font.size = Pt(16)
    
    doc.add_paragraph("\n") 
    
    # 1 a 4
    doc.add_heading("1. Articulación con el plan de desarrollo", level=1)
    doc.add_paragraph(str(plan_desarrollo))
    doc.add_heading("2. Resumen del proyecto", level=1)
    doc.add_paragraph(str(texto_resumen))
    doc.add_heading("3. Marco normativo", level=1)
    doc.add_paragraph(str(texto_normativo))
    doc.add_heading("4. Justificación", level=1)
    doc.add_paragraph(str(justificacion))
    
    # --- 5. LOCALIZACIÓN ---
    doc.add_heading("5. Localización del proyecto", level=1)
    
    if ruta_mapa:
        descargar_y_pegar_imagen(doc, ruta_mapa, 5.0)
    
    doc.add_heading("5.1 Localización", level=2)
    t_loc = doc.add_table(rows=2, cols=6)
    t_loc.style = 'Table Grid'
    headers_loc = ["Departamento", "Provincia", "Municipio", "Barrio/Vereda", "Latitud", "Longitud"]
    
    for i, header_text in enumerate(headers_loc):
        cell = t_loc.cell(0, i)
        cell.text = header_text
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), 'D9E2F3') 
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                
    t_loc.cell(1, 0).text = str(zona_data.get('departamento', ''))
    t_loc.cell(1, 1).text = str(zona_data.get('provincia', ''))
    t_loc.cell(1, 2).text = str(zona_data.get('municipio', ''))
    t_loc.cell(1, 3).text = str(zona_data.get('barrio_vereda', ''))
    t_loc.cell(1, 4).text = str(zona_data.get('latitud', ''))
    t_loc.cell(1, 5).text = str(zona_data.get('longitud', ''))
    
    doc.add_paragraph("\n")
    
    doc.add_heading("5.2 Definición de límites", level=2)
    doc.add_paragraph("Límites Geográficos:", style='List Bullet')
    doc.add_paragraph(str(zona_data.get('limites_geograficos', '')))
    doc.add_paragraph("Límites Administrativos:", style='List Bullet')
    doc.add_paragraph(str(zona_data.get('limites_administrativos', '')))
    doc.add_paragraph("Otros Límites:", style='List Bullet')
    doc.add_paragraph(str(zona_data.get('otros_limites', '')))

    doc.add_heading("5.3 Condiciones de accesibilidad", level=2)
    doc.add_paragraph(str(zona_data.get('accesibilidad', '')))
            
    # --- 6. IDENTIFICACIÓN Y DESCRIPCIÓN DEL PROBLEMA ---
    doc.add_heading("6. Identificación y descripción del problema", level=1)
    
    # 1. Intentar imagen en memoria, 2. Si falla, redibujar
    imagen_insertada = False
    arbol_img_memoria = st.session_state.get('arbol_problemas_img', None) 
    if arbol_img_memoria is not None:
        try:
            p_arbol = doc.add_paragraph()
            p_arbol.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arbol.add_run().add_picture(io.BytesIO(arbol_img_memoria.getvalue()), width=Inches(6.0))
            imagen_insertada = True
        except: pass
        
    if not imagen_insertada and isinstance(datos_h8, dict) and datos_h8:
        img_recreada = redibujar_arbol_problemas(datos_h8)
        if img_recreada:
            p_arbol = doc.add_paragraph()
            p_arbol.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arbol.add_run().add_picture(img_recreada, width=Inches(6.0))
            
    doc.add_heading("6.1 Magnitud del problema", level=2)
    agregar_tabla_word(doc, df_magnitud_reconstruida)
    
    doc.add_paragraph("\n")
    doc.add_heading("6.2 Descripción detallada (Problema - Causa - Efecto)", level=2)
    doc.add_paragraph(str(narrativa_problema))
    
    doc.add_heading("6.3 Antecedentes: ¿Qué se ha hecho previamente con el problema?", level=2)
    doc.add_paragraph(str(antecedentes_problema))
    
    # Registro Fotográfico al final de la sección 6
    if ruta_foto1 or ruta_foto2:
        doc.add_heading("Registro Fotográfico del Problema", level=3)
        if ruta_foto1:
            descargar_y_pegar_imagen(doc, ruta_foto1, 4.5)
        if ruta_foto2:
            descargar_y_pegar_imagen(doc, ruta_foto2, 4.5)

    # --- RESTO DEL DOCUMENTO (7, 8, 9, 10) ---
    doc.add_heading("7. Población", level=1)
    agregar_tabla_word(doc, df_poblacion_general)
    doc.add_heading("1. Población objetivo por sexo", level=2)
    agregar_tabla_word(doc, df_pob_sexo)
    doc.add_heading("2. Población objetivo por rango de edad", level=2)
    agregar_tabla_word(doc, df_pob_edad)
    doc.add_heading("Análisis de la población objetivo", level=2)
    doc.add_paragraph(str(analisis_poblacion))

    doc.add_heading("8. Análisis de Participantes", level=1)
    doc.add_heading("Matriz de Interesados", level=2)
    agregar_tabla_word(doc, df_matriz_interesados)
    doc.add_heading("Mapa de Influencia Estratégico", level=2)
    agregar_tabla_word(doc, df_mapa_influencia)
    doc.add_heading("Análisis de Participantes", level=2)
    agregar_tabla_word(doc, df_analisis_participantes)

    doc.add_heading("9. Objetivos y Resultados Esperados", level=1)
    if arbol_objetivos_img is not None:
        try:
            p_arbol_obj = doc.add_paragraph()
            p_arbol_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arbol_obj.add_run().add_picture(io.BytesIO(arbol_objetivos_img.getvalue()), width=Inches(6.0))
        except:
            pass
            
    doc.add_heading("9.1 Objetivo General", level=2)
    doc.add_paragraph(str(objetivo_general))
    doc.add_heading("9.2 Objetivos Específicos", level=2)
    doc.add_paragraph(str(objetivos_especificos))

    doc.add_heading("10. Alternativas", level=1)
    doc.add_heading("Alternativas Consolidadas", level=2)
    doc.add_paragraph(str(alternativas_consolidadas))
    doc.add_heading("Evaluación de Alternativas", level=2)
    agregar_tabla_word(doc, df_evaluacion_alt)
    doc.add_heading("Alternativa Seleccionada", level=2)
    t_verde = doc.add_table(rows=1, cols=1)
    celda = t_verde.cell(0, 0)
    celda.text = str(alternativa_seleccionada)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'E2F0D9') 
    celda._tc.get_or_add_tcPr().append(shading_elm)
    for paragraph in celda.paragraphs:
        for run in paragraph.runs:
            run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "Reporte en PDF (Aún en construcción)", align="C", new_x="LMARGIN", new_y="NEXT")
    return pdf.output()

# --- BOTONES DE DESCARGA ---
st.markdown('<div class="header-tabla">📥 3. Generar Documento</div>', unsafe_allow_html=True)

col_btn1, col_btn2 = st.columns(2)
with col_btn1:
    st.download_button("📝 Descargar Word (.docx)", data=generar_word(), file_name="Proyecto_Formulado.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
with col_btn2:
    st.download_button("📄 Descargar PDF (.pdf)", data=bytes(generar_pdf()), file_name="Reporte_Prueba.pdf", mime="application/pdf", type="primary", use_container_width=True)
