import streamlit as st
import os
import io
import textwrap
import pandas as pd
import requests
from session_state import inicializar_session, conectar_db, guardar_datos_nube

# --- IMPORTACIÓN DE LIBRERÍAS ---
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    st.error("⚠️ Falta la librería para Word. Agrega 'python-docx' a tu requirements.txt")
    st.stop()

try:
    from fpdf import FPDF
except ImportError:
    st.error("⚠️ Falta la librería para PDF. Agrega 'fpdf2' a tu requirements.txt")
    st.stop()

try:
    import graphviz
except ImportError:
    st.error("⚠️ Falta la librería Graphviz. Agrega 'graphviz' a tu requirements.txt")
    st.stop()

# 1. Asegurar persistencia
inicializar_session()

# ==========================================
# ✅ Persistencia de imágenes de portada (Hoja 16)
# Replica patrón Hoja 9: Storage + (ruta_*, path_*) + guardado nube
# (sin tocar UI)
# ==========================================
if "datos_reportes" not in st.session_state or not isinstance(st.session_state.get("datos_reportes"), dict):
    st.session_state["datos_reportes"] = {}

datos_reportes = st.session_state["datos_reportes"]
datos_reportes.setdefault("ruta_logo_portada", None)
datos_reportes.setdefault("path_logo_portada", None)
datos_reportes.setdefault("sig_logo_portada", None)
datos_reportes.setdefault("ruta_img_portada", None)
datos_reportes.setdefault("path_img_portada", None)
datos_reportes.setdefault("sig_img_portada", None)

# ✅ Persistencia de textos Hoja 16 (6 cajas) - sin cambiar UI
datos_reportes.setdefault("entidad_formulo", "")
datos_reportes.setdefault("division", "")
datos_reportes.setdefault("lugar_presentacion", "Tunja, Boyacá")
datos_reportes.setdefault("anio_presentacion", "2026")
datos_reportes.setdefault("texto_resumen", "")
datos_reportes.setdefault("texto_normativo", "")

def _sync_reportes_field(field_name: str, widget_key: str):
    """Guarda el valor del widget en datos_reportes y lo envía a nube (patrón Hoja 9)."""
    try:
        datos_reportes[field_name] = st.session_state.get(widget_key, "")
        guardar_datos_nube()
    except Exception:
        pass

# Precarga de valores (solo si la clave aún no existe) para no reasignar claves de widgets después de render
if "rep_entidad_formulo" not in st.session_state:
    st.session_state["rep_entidad_formulo"] = datos_reportes.get("entidad_formulo", "")
if "rep_division" not in st.session_state:
    st.session_state["rep_division"] = datos_reportes.get("division", "")
if "rep_lugar_presentacion" not in st.session_state:
    st.session_state["rep_lugar_presentacion"] = datos_reportes.get("lugar_presentacion", "Tunja, Boyacá")
if "rep_anio_presentacion" not in st.session_state:
    st.session_state["rep_anio_presentacion"] = datos_reportes.get("anio_presentacion", "2026")
if "texto_resumen" not in st.session_state:
    st.session_state["texto_resumen"] = datos_reportes.get("texto_resumen", "")
if "texto_normativo" not in st.session_state:
    st.session_state["texto_normativo"] = datos_reportes.get("texto_normativo", "")


def _get_bucket_name() -> str:
    return st.secrets.get("SUPABASE_BUCKET", "uploads")


def _upload_to_supabase_storage_reportes(uploaded_file, tipo_key: str):
    """
    Sube imagen a Supabase Storage y guarda en st.session_state['datos_reportes']:
      - ruta_{tipo_key}: URL pública
      - path_{tipo_key}: path en storage
    Replica patrón de Hoja 9.
    """
    user_id = st.session_state.get("usuario_id")
    if not user_id or uploaded_file is None:
        return None, None

    signature = f"{getattr(uploaded_file, 'name', '')}:{getattr(uploaded_file, 'size', '')}"
    sig_key = f"sig_{tipo_key}"

    # Evita re-subir si es el mismo archivo ya guardado
    if (
        datos_reportes.get(sig_key) == signature
        and datos_reportes.get(f"ruta_{tipo_key}")
        and datos_reportes.get(f"path_{tipo_key}")
    ):
        return datos_reportes.get(f"ruta_{tipo_key}"), datos_reportes.get(f"path_{tipo_key}")

    try:
        db = conectar_db()
        bucket = _get_bucket_name()

        original_name = getattr(uploaded_file, "name", "") or "archivo"
        ext = original_name.split(".")[-1].lower() if "." in original_name else "png"
        content_type = getattr(uploaded_file, "type", None) or "image/png"

        import uuid
        storage_path = f"{user_id}/reportes/{tipo_key}/{uuid.uuid4().hex}.{ext}"

        uploaded_file.seek(0)
        file_bytes = uploaded_file.getvalue()

        db.storage.from_(bucket).upload(
            storage_path,
            file_bytes,
            file_options={"content-type": content_type, "upsert": "true"},
        )

        public_url = db.storage.from_(bucket).get_public_url(storage_path)
        if isinstance(public_url, dict):
            public_url = public_url.get("publicUrl") or public_url.get("public_url")

        datos_reportes[f"ruta_{tipo_key}"] = public_url
        datos_reportes[f"path_{tipo_key}"] = storage_path
        datos_reportes[sig_key] = signature

        # Persistir a nube (el paquete de session_state.py se ajusta después; aquí dejamos el llamado consistente)
        try:
            guardar_datos_nube()
        except Exception:
            pass

        return public_url, storage_path
    except Exception:
        return None, None


def _download_image_bytes(url: str):
    if not url:
        return None
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200:
            return resp.content
    except Exception:
        pass
    return None


# --- DISEÑO PROFESIONAL (CSS) ---
st.markdown("""
    <style>
    .block-container { padding-bottom: 12rem !important; }
    .titulo-seccion { font-size: 30px !important; font-weight: 800 !important; color: #1E3A8A; margin-bottom: 5px; }
    .subtitulo-gris { font-size: 16px !important; color: #666; margin-bottom: 15px; }
    .header-tabla { font-weight: 800; color: #1E3A8A; margin-bottom: 10px; font-size: 1.1rem; text-transform: uppercase; border-bottom: 2px solid #1E3A8A; padding-bottom: 5px;}
    .readonly-box { border: 1px solid #d1d5db; border-radius: 8px; padding: 12px; background-color: #f3f4f6; color: #1E3A8A; font-weight: 800; text-align: center; font-size: 1.2rem;}
    .readonly-autores { border: 1px solid #d1d5db; border-radius: 8px; padding: 10px; background-color: #f3f4f6; color: #374151; font-weight: 600; text-align: center; font-size: 1rem;}
    </style>
""", unsafe_allow_html=True)

# --- ENCABEZADO ---
col_t, col_img = st.columns([4, 1], vertical_alignment="center")
with col_t:
    st.markdown('<div class="titulo-seccion">📄 16. Generador de Reportes</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitulo-gris">Configuración de portada y exportación del documento final.</div>', unsafe_allow_html=True)
with col_img:
    if os.path.exists("unnamed.jpg"):
        st.image("unnamed.jpg", use_container_width=True)

st.divider()

# ==========================================
# 📘 1. CONFIGURACIÓN DE LA PORTADA
# ==========================================
st.markdown('<div class="header-tabla">📘 1. Configuración de la Portada</div>', unsafe_allow_html=True)

nombre_proyecto = st.session_state.get('nombre_proyecto_libre', 'NOMBRE DEL PROYECTO NO DEFINIDO')

st.write("**Nombre del Proyecto:**")
st.markdown(f'<div class="readonly-box">{nombre_proyecto.upper()}</div><br>', unsafe_allow_html=True)

# --- LIMPIEZA VISUAL DE LAS IMÁGENES ---
col_up1, col_up2 = st.columns(2)
with col_up1:
    logo_entidad = st.file_uploader("🖼️ Sube el Logo de la Entidad", type=["png", "jpg", "jpeg"], key="logo_portada")
with col_up2:
    img_portada = st.file_uploader("📸 Sube la Imagen Central", type=["png", "jpg", "jpeg"], key="img_portada")

ruta_logo_guardada = datos_reportes.get("ruta_logo_portada")
ruta_img_guardada = datos_reportes.get("ruta_img_portada")

if (logo_entidad is not None) or (img_portada is not None) or ruta_logo_guardada or ruta_img_guardada:
    col_prev1, col_prev2 = st.columns(2)
    with col_prev1:
        if logo_entidad is not None:
            st.image(logo_entidad, width=150)
        elif ruta_logo_guardada:
            st.image(ruta_logo_guardada, width=150)
    with col_prev2:
        if img_portada is not None:
            st.image(img_portada, width=300)
        elif ruta_img_guardada:
            st.image(ruta_img_guardada, width=300)

# ✅ Guardado en Storage (solo lógica; UI intacta)
if logo_entidad is not None:
    _upload_to_supabase_storage_reportes(logo_entidad, "logo_portada")
if img_portada is not None:
    _upload_to_supabase_storage_reportes(img_portada, "img_portada")

st.write("")

nombres_formuladores = "No se encontraron formuladores registrados en la Hoja 0 (Equipo)"
nombres_display = nombres_formuladores

integrantes = st.session_state.get("integrantes", [])
if isinstance(integrantes, list):
    nombres_validos = []
    for p in integrantes:
        if isinstance(p, dict):
            nombre = str(p.get("Nombre Completo", "")).strip()
            if nombre:
                nombres_validos.append(nombre)
    if nombres_validos:
        nombres_formuladores = "\n".join(nombres_validos)
        nombres_display = ", ".join(nombres_validos)

st.write("**Presentado por (Equipo Formulador):**")
st.markdown(f'<div class="readonly-autores">{nombres_display if "nombres_display" in locals() else nombres_formuladores}</div><br>', unsafe_allow_html=True)

col_d1, col_d2 = st.columns(2)
with col_d1:
    entidad_formulo = st.text_input("Entidad que formula el proyecto", placeholder="Ej: Alcaldía de Tunja", key="rep_entidad_formulo", on_change=_sync_reportes_field, args=("entidad_formulo", "rep_entidad_formulo"))
with col_d2:
    division = st.text_input("División / Dependencia", placeholder="Ej: Secretaría de Infraestructura", key="rep_division", on_change=_sync_reportes_field, args=("division", "rep_division"))

col_d3, col_d4 = st.columns(2)
with col_d3:
    lugar_presentacion = st.text_input("Lugar de presentación", key="rep_lugar_presentacion", on_change=_sync_reportes_field, args=("lugar_presentacion", "rep_lugar_presentacion"))
with col_d4:
    anio_presentacion = st.text_input("Año", key="rep_anio_presentacion", on_change=_sync_reportes_field, args=("anio_presentacion", "rep_anio_presentacion"))

st.divider()

# ==========================================
# 📑 2. SELECCIÓN Y DILIGENCIAMIENTO DE CONTENIDO
# ==========================================
# --- NUEVO TÍTULO Y CAJAS AUTOAJUSTABLES ---
st.markdown('<div class="header-tabla">📑 2. Resumen y Marco Normativo</div>', unsafe_allow_html=True)
st.write("Diligencia las siguientes secciones que se incluirán al inicio de tu documento:")

instrucciones_resumen = "El resumen es el elemento fundamental para dar contexto sobre el proyecto, en este sentido escriba en máximo 15 líneas de manera clara, sencilla, directa y concisa el resumen del contenido del proyecto, que permitan dar una idea sobre el alcance, componentes y productos esperados."

curr_resumen = st.session_state.get("texto_resumen", "")
h_res = int(max(150, (curr_resumen.count('\n') + (len(curr_resumen) / 100) + 1) * 25 + 40))
texto_resumen = st.text_area("2. RESUMEN DEL PROYECTO", placeholder=instrucciones_resumen, height=h_res, key="texto_resumen", on_change=_sync_reportes_field, args=("texto_resumen", "texto_resumen"))

curr_norm = st.session_state.get("texto_normativo", "")
h_norm = int(max(150, (curr_norm.count('\n') + (len(curr_norm) / 100) + 1) * 25 + 40))
texto_normativo = st.text_area("3. MARCO NORMATIVO", placeholder="Escriba aquí el marco normativo del proyecto...", height=h_norm, key="texto_normativo", on_change=_sync_reportes_field, args=("texto_normativo", "texto_normativo"))

st.divider()

# ==========================================
# 📥 EXTRACCIÓN DE DATOS DE LA MEMORIA
# ==========================================
# 1 a 4. General
_plan_nombre = str(st.session_state.get('plan_nombre', '')).strip()
_plan_eje = str(st.session_state.get('plan_eje', '')).strip()
_plan_programa = str(st.session_state.get('plan_programa', '')).strip()

if any([_plan_nombre, _plan_eje, _plan_programa]):
    plan_desarrollo = (
        f"Nombre del Plan: {_plan_nombre if _plan_nombre else 'No definido'}\n"
        f"Eje: {_plan_eje if _plan_eje else 'No definido'}\n"
        f"Programa: {_plan_programa if _plan_programa else 'No definido'}"
    )
else:
    plan_desarrollo = st.session_state.get('plan_desarrollo', 'No se ha registrado información en la Hoja 15.')

justificacion = (
    st.session_state.get('justificacion_arbol_objetivos_final')
    or st.session_state.get('arbol_objetivos_final', {}).get('referencia_manual', {}).get('justificacion', '')
    or st.session_state.get('justificacion_proyecto')
    or 'No se ha registrado información en la Hoja 7.'
)

# 5. Localización
zona_data = st.session_state.get('descripcion_zona', {})
ruta_mapa = zona_data.get('ruta_mapa')
ruta_foto1 = zona_data.get('ruta_foto1')
ruta_foto2 = zona_data.get('ruta_foto2')

# --- LÓGICA DE EXTRACCIÓN PARA LA SECCIÓN 6 ---
def _a_texto_dict(item):
    if isinstance(item, dict): return item
    if item is None: return None
    return {"texto": str(item)}

def _a_lista_dicts(valor):
    if valor is None: return []
    if isinstance(valor, list):
        out = []
        for it in valor:
            d = _a_texto_dict(it)
            if d is not None: out.append(d)
        return out
    if isinstance(valor, dict): return [valor]
    return [{"texto": str(valor)}]

desc_prob_data = st.session_state.get('descripcion_problema', {})
narrativa_problema = desc_prob_data.get('redaccion_narrativa', 'No se ha registrado descripción.')
if not narrativa_problema.strip(): narrativa_problema = 'No se ha registrado descripción.'

antecedentes_problema = desc_prob_data.get('antecedentes', 'No se han registrado antecedentes.')
if not antecedentes_problema.strip(): antecedentes_problema = 'No se han registrado antecedentes.'

tabla_datos_prob = desc_prob_data.get('tabla_datos', {})
datos_h8 = st.session_state.get('arbol_problemas_final', {})
if not isinstance(datos_h8, dict): datos_h8 = {}

pp_list = _a_lista_dicts(datos_h8.get("Problema Principal"))
pc_txt = pp_list[0].get("texto", "") if pp_list and isinstance(pp_list[0], dict) else ""

lista_causas = [c.get("texto") for c in (_a_lista_dicts(datos_h8.get("Causas Directas")) + _a_lista_dicts(datos_h8.get("Causas Indirectas"))) if isinstance(c, dict) and c.get("texto")]
lista_efectos = [e.get("texto") for e in (_a_lista_dicts(datos_h8.get("Efectos Directos")) + _a_lista_dicts(datos_h8.get("Efectos Indirectos"))) if isinstance(e, dict) and e.get("texto")]

filas_magnitud = []
if pc_txt:
    filas_magnitud.append({"Categoría": "PROBLEMA CENTRAL", "Descripción": pc_txt, "Magnitud": tabla_datos_prob.get("m_pc", ""), "Unidad": tabla_datos_prob.get("u_pc", ""), "Cantidad": tabla_datos_prob.get("c_pc", "")})
for i, txt in enumerate(lista_causas):
    filas_magnitud.append({"Categoría": f"CAUSA {i+1}", "Descripción": txt, "Magnitud": tabla_datos_prob.get(f"m_causa_{i}", ""), "Unidad": tabla_datos_prob.get(f"u_causa_{i}", ""), "Cantidad": tabla_datos_prob.get(f"c_causa_{i}", "")})
for i, txt in enumerate(lista_efectos):
    filas_magnitud.append({"Categoría": f"EFECTO {i+1}", "Descripción": txt, "Magnitud": tabla_datos_prob.get(f"m_efecto_{i}", ""), "Unidad": tabla_datos_prob.get(f"u_efecto_{i}", ""), "Cantidad": tabla_datos_prob.get(f"c_efecto_{i}", "")})

df_magnitud_reconstruida = pd.DataFrame(filas_magnitud)

# --- 7. POBLACIÓN ---
df_poblacion_general = pd.DataFrame([{
    "Población de Referencia": zona_data.get('poblacion_referencia', 0),
    "Población Afectada": zona_data.get('poblacion_afectada', 0),
    "Población Objetivo": zona_data.get('poblacion_objetivo', 0)
}])

genero_data = zona_data.get('poblacion_objetivo_genero', {})
df_pob_sexo = pd.DataFrame([{
    "Hombres": genero_data.get("Hombres", 0),
    "Mujeres": genero_data.get("Mujeres", 0)
}])

edad_data = zona_data.get('poblacion_objetivo_edad', {})
df_pob_edad = pd.DataFrame([{
    "0 - 14": edad_data.get("0-14", 0),
    "15 - 19": edad_data.get("15-19", 0),
    "20 - 59": edad_data.get("20-59", 0),
    "Mayor de 60 años": edad_data.get("Mayor de 60 años", 0)
}])

analisis_poblacion = str(zona_data.get('analisis_poblacion_objetivo', '')).strip()
if not analisis_poblacion:
    analisis_poblacion = 'No se ha registrado análisis de población.'

# --- 8. PARTICIPANTES ---
df_matriz_interesados = st.session_state.get('df_interesados', pd.DataFrame())
if df_matriz_interesados is not None and not isinstance(df_matriz_interesados, pd.DataFrame):
    try: df_matriz_interesados = pd.DataFrame(df_matriz_interesados)
    except: df_matriz_interesados = pd.DataFrame()

texto_analisis_participantes = str(
    st.session_state.get('analisis_participantes') or
    st.session_state.get('txt_analisis_participantes') or
    st.session_state.get('descripcion_zona', {}).get('analisis_participantes', '')
).strip()

if not texto_analisis_participantes:
    texto_analisis_participantes = "No se ha registrado el análisis de los participantes."

# --- 9. OBJETIVOS ---
arbol_obj_datos = st.session_state.get('arbol_objetivos_final', {})
ref_obj_data = arbol_obj_datos.get('referencia_manual', {})

objetivo_general = str(ref_obj_data.get('objetivo', '')).strip()
if not objetivo_general:
    objetivo_general = "No se ha definido el objetivo general en la Hoja 7."

objetivos_especificos = ref_obj_data.get('especificos', [])
if not objetivos_especificos:
    objetivos_especificos = ["No se han definido objetivos específicos."]

# --- 10. ALTERNATIVAS ---
lista_alts_evaluadas = st.session_state.get('lista_alternativas', [])
pesos = st.session_state.get('ponderacion_criterios', {"COSTO": 25, "FACILIDAD": 25, "BENEFICIOS": 25, "TIEMPO": 25})
df_criterios = pd.DataFrame([pesos])

df_calif = st.session_state.get('df_calificaciones', pd.DataFrame())
df_evaluacion_alt = pd.DataFrame()
alternativa_seleccionada = "No se ha seleccionado ninguna alternativa."

if not df_calif.empty:
    df_eval = df_calif.copy()
    df_eval["TOTAL"] = 0.0
    for c in ["COSTO", "FACILIDAD", "BENEFICIOS", "TIEMPO"]:
        if c in df_eval.columns:
            df_eval["TOTAL"] += df_eval[c] * (pesos.get(c, 0) / 100.0)

    ganadora_idx = df_eval["TOTAL"].idxmax()
    puntaje_ganador = df_eval["TOTAL"].max()
    alternativa_seleccionada = f"Alternativa Seleccionada: {ganadora_idx} ({puntaje_ganador:.2f} pts)"

    df_eval = df_eval.round(2).reset_index().rename(columns={"index": "Alternativa"})
    df_evaluacion_alt = df_eval

# ==========================================
# ⚙️ MOTORES DE DIBUJO Y GENERACIÓN WORD
# ==========================================
def agregar_tabla_word(doc, df):
    if isinstance(df, pd.DataFrame) and not df.empty:
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Table Grid'
        hdr_cells = t.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for index, row in df.iterrows():
            row_cells = t.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
    else:
        p = doc.add_paragraph()
        p.add_run("No se registraron datos en esta tabla.").italic = True

def descargar_y_pegar_imagen(doc, url, ancho):
    if url:
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(io.BytesIO(response.content), width=Inches(ancho))
                return True
        except Exception as e:
            pass
    return False

def redibujar_arbol_problemas(arbol_data):
    try:
        dot = graphviz.Digraph(format='png')
        dot.attr(rankdir='BT')
        dot.attr('node', shape='box', style='filled', fontname='Helvetica', margin='0.2')
        pp_list = _a_lista_dicts(arbol_data.get("Problema Principal", arbol_data.get("problema")))
        pc_txt = pp_list[0].get("texto", "Problema Central") if pp_list else "Problema Central"
        dot.node('PC', str(pc_txt), fillcolor='#FCA5A5')
        c_dir = _a_lista_dicts(arbol_data.get("Causas Directas", arbol_data.get("causas")))
        c_ind = _a_lista_dicts(arbol_data.get("Causas Indirectas", []))
        causas = c_dir + c_ind
        for i, ca in enumerate(causas):
            if ca.get('texto'):
                dot.node(f'C_{i}', ca.get('texto'), fillcolor='#FEF3C7')
                dot.edge(f'C_{i}', 'PC')
        e_dir = _a_lista_dicts(arbol_data.get("Efectos Directos", arbol_data.get("efectos")))
        e_ind = _a_lista_dicts(arbol_data.get("Efectos Indirectos", []))
        efectos = e_dir + e_ind
        for i, ef in enumerate(efectos):
            if ef.get('texto'):
                dot.node(f'E_{i}', ef.get('texto'), fillcolor='#DBEAFE')
                dot.edge('PC', f'E_{i}')
        return io.BytesIO(dot.pipe())
    except Exception as e:
        return None

def redibujar_arbol_objetivos(datos):
    try:
        CONFIG_OBJ = {
            "Fin Último":        {"color": "#0E6251"},
            "Fines Indirectos":  {"color": "#154360"},
            "Fines Directos":    {"color": "#1F618D"},
            "Objetivo General":  {"color": "#C0392B"},
            "Medios Directos":   {"color": "#F1C40F"},
            "Medios Indirectos": {"color": "#D35400"}
        }
        claves_graficas = [k for k in datos.keys() if k != 'referencia_manual']
        if not any(datos.get(k) for k in claves_graficas): return None
        dot = graphviz.Digraph(format='png')
        dot.attr(rankdir='BT', nodesep='0.4', ranksep='0.6', splines='ortho')
        dot.attr('node', fontsize='11', fontname='Arial', style='filled', shape='box', margin='0.3,0.2', width='2.5')
        def limpiar(t): return "\\n".join(textwrap.wrap(str(t).upper(), width=25))
        obj_gen = [it for it in datos.get("Objetivo General", []) if isinstance(it, dict) and it.get('texto')]
        if obj_gen:
            dot.node("OG", limpiar(obj_gen[0]['texto']), fillcolor=CONFIG_OBJ["Objetivo General"]["color"], fontcolor='white', color='none', width='4.5')
        for tipo, p_id, h_tipo in [("Fines Directos", "OG", "Fines Indirectos"), ("Medios Directos", "OG", "Medios Indirectos")]:
            items = [it for it in datos.get(tipo, []) if isinstance(it, dict) and it.get('texto')]
            for i, item in enumerate(items):
                n_id = f"{tipo[:2]}{i}"
                dot.node(n_id, limpiar(item['texto']), fillcolor=CONFIG_OBJ[tipo]["color"], fontcolor='white' if "Fin" in tipo else 'black', color='none')
                if "Fin" in tipo: dot.edge("OG", n_id)
                else: dot.edge(n_id, "OG")
                hijos = [h for h in datos.get(h_tipo, []) if isinstance(h, dict) and h.get('padre') == item.get('texto')]
                for j, h in enumerate(hijos):
                    h_id = f"{h_tipo[:2]}{i}_{j}"
                    dot.node(h_id, limpiar(h['texto']), fillcolor=CONFIG_OBJ[h_tipo]["color"], fontcolor='white', color='none', fontsize='10')
                    if "Fin" in tipo: dot.edge(n_id, h_id)
                    else: dot.edge(h_id, n_id)
        return io.BytesIO(dot.pipe())
    except Exception as e:
        return None

def generar_word():
    doc = Document()

    # --- UNIFICACIÓN DE COLOR DE TÍTULOS (AZUL OSCURO #1E3A8A) ---
    color_azul_oscuro = RGBColor(30, 58, 138)
    for i in range(1, 5):
        try:
            style = doc.styles[f'Heading {i}']
            style.font.color.rgb = color_azul_oscuro
            style.font.bold = True
        except:
            pass

    try:
        doc.styles['Title'].font.color.rgb = color_azul_oscuro
    except:
        pass

    # --- CONFIGURACIÓN DE ENCABEZADO Y PIE DE PÁGINA ---
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    def crear_encabezado(hdr_obj):
        htable = hdr_obj.add_table(rows=1, cols=2, width=Inches(6.5))
        htable.autofit = False
        htable.columns[0].width = Inches(5.0)
        htable.columns[1].width = Inches(1.5)
        h_izq = htable.cell(0, 0).paragraphs[0]
        h_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_izq.add_run(nombre_proyecto.upper()).bold = True
        h_der = htable.cell(0, 1).paragraphs[0]
        h_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ✅ Si no hay archivo en la sesión, intenta usar el guardado en Storage
        if logo_entidad is not None:
            logo_entidad.seek(0)
            h_der.add_run().add_picture(io.BytesIO(logo_entidad.getvalue()), width=Inches(0.6))
        else:
            url_logo = datos_reportes.get("ruta_logo_portada")
            b = _download_image_bytes(url_logo)
            if b:
                h_der.add_run().add_picture(io.BytesIO(b), width=Inches(0.6))

        p_line = hdr_obj.add_paragraph()
        pPr = p_line._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    crear_encabezado(section.first_page_header)
    crear_encabezado(section.header)

    footer = section.footer
    p_line_f = footer.paragraphs[0]
    pPr_f = p_line_f._p.get_or_add_pPr()
    pBdr_f = OxmlElement('w:pBdr')
    bottom_f = OxmlElement('w:bottom')
    bottom_f.set(qn('w:val'), 'single')
    bottom_f.set(qn('w:sz'), '6')
    bottom_f.set(qn('w:space'), '1')
    bottom_f.set(qn('w:color'), 'auto')
    pBdr_f.append(bottom_f)
    pPr_f.append(pBdr_f)

    ftable = footer.add_table(rows=1, cols=3, width=Inches(7.0))
    ftable.autofit = False

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    ftable._tbl.tblPr.append(tblLayout)

    ancho_izq = Inches(0.5)
    ancho_cen = Inches(6.0)
    ancho_der = Inches(0.5)

    ftable.columns[0].width = ancho_izq
    ftable.columns[1].width = ancho_cen
    ftable.columns[2].width = ancho_der
    for row in ftable.rows:
        row.cells[0].width = ancho_izq
        row.cells[1].width = ancho_cen
        row.cells[2].width = ancho_der

    f_centro = ftable.cell(0, 1).paragraphs[0]
    f_centro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_centro.paragraph_format.space_after = Pt(0)
    f_centro.paragraph_format.space_before = Pt(0)

    texto_entidad = entidad_formulo if entidad_formulo else ""
    texto_div = division if division else ""

    if texto_entidad:
        r_centro1 = f_centro.add_run(texto_entidad.upper())
        r_centro1.font.size = Pt(9)
        r_centro1.italic = True
        r_centro1.font.color.rgb = RGBColor(128, 128, 128)

    if texto_div:
        if texto_entidad:
            f_centro2 = ftable.cell(0, 1).add_paragraph()
        else:
            f_centro2 = f_centro

        f_centro2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_centro2.paragraph_format.space_after = Pt(0)
        f_centro2.paragraph_format.space_before = Pt(0)

        r_centro2 = f_centro2.add_run(texto_div.upper())
        r_centro2.font.size = Pt(9)
        r_centro2.italic = True
        r_centro2.font.color.rgb = RGBColor(128, 128, 128)

    f_der = ftable.cell(0, 2).paragraphs[0]
    f_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_num = f_der.add_run()
    r_num.font.size = Pt(10)
    r_num.font.color.rgb = RGBColor(128, 128, 128)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_num._r.append(fldChar1)
    r_num._r.append(instrText)
    r_num._r.append(fldChar2)
    r_num._r.append(fldChar3)

    # --- 1. PORTADA (AJUSTADA SIN EXCESO DE SALTOS DE LÍNEA) ---
    doc.add_paragraph()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titulo = p_titulo.add_run(nombre_proyecto.upper())
    r_titulo.bold = True
    r_titulo.font.size = Pt(20)

    doc.add_paragraph()
    if img_portada is not None:
        img_portada.seek(0)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(img_portada.getvalue()), width=Inches(3.8))
    else:
        url_img = datos_reportes.get("ruta_img_portada")
        b = _download_image_bytes(url_img)
        if b:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(io.BytesIO(b), width=Inches(3.8))

    doc.add_paragraph()
    if entidad_formulo:
        doc.add_paragraph(entidad_formulo.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if division:
        doc.add_paragraph(division.upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p_presentado = doc.add_paragraph()
    p_presentado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_presentado.add_run("Presentado por:\n").italic = True
    p_presentado.add_run(nombres_formuladores).bold = True

    doc.add_paragraph()
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto_lugar = lugar_presentacion if lugar_presentacion else ""
    texto_anio = anio_presentacion if anio_presentacion else ""
    p_pie.add_run(f"{texto_lugar}\n{texto_anio}".strip()).bold = True

    doc.add_page_break()

    # --- 2. INICIO DEL CONTENIDO ---
    p_tit_cont = doc.add_paragraph()
    p_tit_cont.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t_cont = p_tit_cont.add_run(nombre_proyecto.upper())
    r_t_cont.bold = True
    r_t_cont.font.size = Pt(16)

    doc.add_paragraph("\n")

    # 1 a 4
    doc.add_heading("1. Articulación con el plan de desarrollo", level=1)
    doc.add_paragraph(str(plan_desarrollo))
    doc.add_heading("2. Resumen del proyecto", level=1)
    doc.add_paragraph(str(texto_resumen))
    doc.add_heading("3. Marco normativo", level=1)
    doc.add_paragraph(str(texto_normativo))
    doc.add_heading("4. Justificación", level=1)
    doc.add_paragraph(str(justificacion))

    # 5. Localización
    doc.add_heading("5. Localización del proyecto", level=1)

    if ruta_mapa:
        descargar_y_pegar_imagen(doc, ruta_mapa, 5.0)

    doc.add_heading("5.1 Localización", level=2)
    t_loc = doc.add_table(rows=2, cols=6)
    t_loc.style = 'Table Grid'
    headers_loc = ["Departamento", "Provincia", "Municipio", "Barrio/Vereda", "Latitud", "Longitud"]

    for i, header_text in enumerate(headers_loc):
        cell = t_loc.cell(0, i)
        cell.text = header_text
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), 'D9E2F3')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    t_loc.cell(1, 0).text = str(zona_data.get('departamento', ''))
    t_loc.cell(1, 1).text = str(zona_data.get('provincia', ''))
    t_loc.cell(1, 2).text = str(zona_data.get('municipio', ''))
    t_loc.cell(1, 3).text = str(zona_data.get('barrio_vereda', ''))
    t_loc.cell(1, 4).text = str(zona_data.get('latitud', ''))
    t_loc.cell(1, 5).text = str(zona_data.get('longitud', ''))

    doc.add_paragraph("\n")

    doc.add_heading("5.2 Definición de límites", level=2)
    doc.add_paragraph("Límites Geográficos:", style='List Bullet')
    doc.add_paragraph(str(zona_data.get('limites_geograficos', '')))
    doc.add_paragraph("Límites Administrativos:", style='List Bullet')
    doc.add_paragraph(str(zona_data.get('limites_administrativos', '')))
    doc.add_paragraph("Otros Límites:", style='List Bullet')
    doc.add_paragraph(str(zona_data.get('otros_limites', '')))

    doc.add_heading("5.3 Condiciones de accesibilidad", level=2)
    doc.add_paragraph(str(zona_data.get('accesibilidad', '')))

    # --- 6. IDENTIFICACIÓN Y DESCRIPCIÓN DEL PROBLEMA ---
    doc.add_heading("6. Identificación y descripción del problema", level=1)

    imagen_prob_insertada = False
    arbol_prob_memoria = st.session_state.get('arbol_problemas_img', None)
    if arbol_prob_memoria is not None:
        try:
            p_arbol_p = doc.add_paragraph()
            p_arbol_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arbol_p.add_run().add_picture(io.BytesIO(arbol_prob_memoria.getvalue()), width=Inches(6.0))
            imagen_prob_insertada = True
        except:
            pass

    if not imagen_prob_insertada and datos_h8:
        img_prob_recreada = redibujar_arbol_problemas(datos_h8)
        if img_prob_recreada:
            p_arbol_p = doc.add_paragraph()
            p_arbol_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arbol_p.add_run().add_picture(img_prob_recreada, width=Inches(6.0))

    doc.add_heading("6.1 Magnitud del problema", level=2)
    agregar_tabla_word(doc, df_magnitud_reconstruida)

    doc.add_paragraph("\n")
    doc.add_heading("Descripción detallada (Problema - Causa - Efecto)", level=3)
    doc.add_paragraph(str(narrativa_problema))

    doc.add_heading("Antecedentes: ¿Qué se ha hecho previamente con el problema?", level=3)
    doc.add_paragraph(str(antecedentes_problema))

    if ruta_foto1 or ruta_foto2:
        doc.add_heading("Registro Fotográfico del Problema", level=3)
        if ruta_foto1:
            descargar_y_pegar_imagen(doc, ruta_foto1, 4.5)
        if ruta_foto2:
            descargar_y_pegar_imagen(doc, ruta_foto2, 4.5)

    # --- 7. POBLACIÓN ---
    doc.add_heading("7. Población", level=1)
    agregar_tabla_word(doc, df_poblacion_general)

    doc.add_heading("7.1 Población objetivo por sexo", level=2)
    agregar_tabla_word(doc, df_pob_sexo)

    doc.add_heading("7.2 Población objetivo por rango de edad", level=2)
    agregar_tabla_word(doc, df_pob_edad)

    doc.add_heading("7.3 Análisis de la población objetivo", level=2)
    doc.add_paragraph(str(analisis_poblacion))

    # --- 8. PARTICIPANTES ---
    doc.add_heading("8. Análisis de Participantes", level=1)
    doc.add_heading("Matriz de Interesados", level=2)
    agregar_tabla_word(doc, df_matriz_interesados)
    doc.add_paragraph("\n" + str(texto_analisis_participantes))

    # --- 9. OBJETIVOS ---
    doc.add_heading("9. Objetivos", level=1)

    imagen_obj_insertada = False
    arbol_obj_memoria = st.session_state.get('arbol_objetivos_img', None)
    if arbol_obj_memoria is not None:
        try:
            p_arbol_obj = doc.add_paragraph()
            p_arbol_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arbol_obj.add_run().add_picture(io.BytesIO(arbol_obj_memoria.getvalue()), width=Inches(6.0))
            imagen_obj_insertada = True
        except:
            pass

  if not imagen_obj_insertada and arbol_obj_datos:
        img_obj_recreada = redibujar_arbol_objetivos(arbol_obj_datos)
        if img_obj_recreada:
            try:
                p_arbol_obj = doc.add_paragraph()
                p_arbol_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # ✅ CORRECCIÓN: Usamos la imagen directamente sin el io.BytesIO() extra
                p_arbol_obj.add_run().add_picture(img_obj_recreada, width=Inches(6.0))
                imagen_obj_insertada = True
            except Exception:
                # Si algo falla con la imagen, el programa sigue adelante sin romperse
                doc.add_paragraph("[Aviso: El Árbol de Objetivos no se pudo renderizar]")

    doc.add_heading("9.1 Objetivo General", level=2)
    doc.add_paragraph(str(objetivo_general))

    doc.add_heading("9.2 Objetivos Específicos", level=2)
    for oe in objetivos_especificos:
        if str(oe).strip():
            doc.add_paragraph(str(oe).strip(), style='List Bullet')

    # --- 10. ALTERNATIVAS ---
    doc.add_heading("10. Alternativas", level=1)

    doc.add_heading("10.1 Alternativas Evaluadas", level=2)
    if not lista_alts_evaluadas:
        doc.add_paragraph("No se han registrado alternativas evaluadas en la Hoja 6.")
    else:
        for i, alt in enumerate(lista_alts_evaluadas):
            p_alt = doc.add_paragraph()
            p_alt.add_run(f"{i+1}. {alt['nombre'].upper()}").bold = True
            if alt.get('descripcion'):
                doc.add_paragraph(f"Descripción: {alt['descripcion']}")
            for conf in alt.get('configuracion', []):
                doc.add_paragraph(f"🎯 Objetivo: {conf['objetivo']}", style='List Bullet')
                for act in conf['actividades']:
                    p_act = doc.add_paragraph(f"🔹 Actividad: {act}")
                    p_act.paragraph_format.left_indent = Inches(0.5)
            doc.add_paragraph("\n")

    doc.add_heading("10.2 Evaluación de Alternativas", level=2)

    doc.add_heading("Criterios Evaluados (Pesos %)", level=3)
    agregar_tabla_word(doc, df_criterios)
    doc.add_paragraph("\n")

    doc.add_heading("Calificaciones y Puntaje Total", level=3)
    agregar_tabla_word(doc, df_evaluacion_alt)
    doc.add_paragraph("\n")

    doc.add_heading("Alternativa Seleccionada", level=3)
    t_verde = doc.add_table(rows=1, cols=1)
    celda = t_verde.cell(0, 0)
    celda.text = str(alternativa_seleccionada)

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), 'E2F0D9')
    celda._tc.get_or_add_tcPr().append(shading_elm)

    for paragraph in celda.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "Reporte en PDF (Aún en construcción)", align="C", new_x="LMARGIN", new_y="NEXT")
    return pdf.output()

# --- BOTONES DE DESCARGA ---
st.markdown('<div class="header-tabla">📥 3. Generar Documento</div>', unsafe_allow_html=True)

col_btn1, col_btn2 = st.columns(2)
with col_btn1:
    st.download_button("📝 Descargar Word (.docx)", data=generar_word(), file_name="Proyecto_Formulado.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
with col_btn2:
    st.download_button("📄 Descargar PDF (.pdf)", data=bytes(generar_pdf()), file_name="Reporte_Prueba.pdf", mime="application/pdf", type="primary", use_container_width=True)
